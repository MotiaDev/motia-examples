"""
Document Parser Service for Airbnb Guest Assistant

Parses various document formats (PDF, Word, HTML, Markdown) into
a normalized structure preserving headings and sections.
"""

import os
import re
import httpx
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass
from enum import Enum
import chardet


class DocType(str, Enum):
    """Document type classification."""
    HOUSE_MANUAL = "house_manual"
    LOCAL_GUIDE = "local_guide"
    APPLIANCE_MANUAL = "appliance_manual"
    POLICY = "policy"
    UNKNOWN = "unknown"


@dataclass
class ParsedSection:
    """A section extracted from a document."""
    title: str
    content: str
    level: int  # Heading level (1-6)
    is_critical: bool  # Safety, rules, emergency info


@dataclass
class ParsedDocument:
    """A fully parsed document."""
    title: str
    sections: List[ParsedSection]
    doc_type: DocType
    language: str
    source_url: str
    source_filename: str
    raw_text: str
    metadata: Dict[str, Any]


class DocumentParser:
    """Service for parsing various document formats."""
    
    # Keywords that indicate critical sections
    CRITICAL_KEYWORDS = [
        "safety", "emergency", "fire", "evacuation", "rules", "house rules",
        "prohibited", "not allowed", "forbidden", "warning", "caution",
        "important", "must", "required", "checkout", "check-out", "check out",
        "checkin", "check-in", "check in", "late checkout", "early checkin",
        "security", "alarm", "code", "lock", "key", "lockbox", "keypad",
        "contact", "host contact", "emergency contact", "911", "police",
        "hospital", "medical", "first aid", "poison", "gas leak"
    ]
    
    # Keywords for doc type classification
    DOC_TYPE_KEYWORDS = {
        DocType.HOUSE_MANUAL: [
            "welcome", "house manual", "property guide", "home guide",
            "your stay", "amenities", "facilities", "room", "bedroom",
            "bathroom", "kitchen", "living"
        ],
        DocType.LOCAL_GUIDE: [
            "local", "neighborhood", "area guide", "restaurant", "cafe",
            "attractions", "transport", "bus", "train", "metro", "subway",
            "airport", "taxi", "uber", "grocery", "supermarket", "shopping",
            "things to do", "nearby", "directions"
        ],
        DocType.APPLIANCE_MANUAL: [
            "appliance", "how to use", "instructions", "operating",
            "washing machine", "dishwasher", "oven", "stove", "microwave",
            "thermostat", "heating", "cooling", "ac", "air conditioning",
            "tv", "television", "wifi", "wi-fi", "internet", "remote"
        ],
        DocType.POLICY: [
            "policy", "terms", "rules", "cancellation", "refund",
            "payment", "deposit", "damage", "cleaning fee", "pet policy",
            "smoking", "noise", "quiet hours", "guest policy", "visitors"
        ]
    }
    
    def __init__(self):
        """Initialize the parser."""
        self.http_client = httpx.AsyncClient(
            timeout=60.0,
            follow_redirects=True,
            headers={"User-Agent": "AirbnbAssistant/1.0"}
        )
    
    async def close(self):
        """Close the HTTP client."""
        await self.http_client.aclose()
    
    async def fetch_url(self, url: str) -> Tuple[bytes, str, Dict[str, str]]:
        """
        Fetch content from a URL.
        
        Args:
            url: The URL to fetch
            
        Returns:
            Tuple of (content_bytes, content_type, headers)
        """
        response = await self.http_client.get(url)
        response.raise_for_status()
        
        content_type = response.headers.get("content-type", "").lower()
        return response.content, content_type, dict(response.headers)
    
    def detect_encoding(self, content: bytes) -> str:
        """Detect text encoding."""
        result = chardet.detect(content)
        return result.get("encoding", "utf-8") or "utf-8"
    
    async def parse_url(
        self, 
        url: str, 
        property_id: str,
        language: str = "en"
    ) -> ParsedDocument:
        """
        Parse a document from URL.
        
        Args:
            url: URL to fetch and parse
            property_id: Property identifier
            language: Document language
            
        Returns:
            ParsedDocument with extracted content
        """
        content, content_type, headers = await self.fetch_url(url)
        
        # Determine format and parse
        if "pdf" in content_type or url.lower().endswith(".pdf"):
            return await self._parse_pdf(content, url, "", property_id, language)
        elif "html" in content_type or url.lower().endswith(".html"):
            encoding = self.detect_encoding(content)
            text = content.decode(encoding, errors="replace")
            return self._parse_html(text, url, "", property_id, language)
        elif any(ext in url.lower() for ext in [".md", ".markdown"]):
            encoding = self.detect_encoding(content)
            text = content.decode(encoding, errors="replace")
            return self._parse_markdown(text, url, "", property_id, language)
        elif any(ext in url.lower() for ext in [".doc", ".docx"]):
            return await self._parse_docx(content, url, "", property_id, language)
        else:
            # Try to parse as HTML or plain text
            encoding = self.detect_encoding(content)
            text = content.decode(encoding, errors="replace")
            if "<html" in text.lower() or "<body" in text.lower():
                return self._parse_html(text, url, "", property_id, language)
            else:
                return self._parse_plain_text(text, url, "", property_id, language)
    
    async def parse_file(
        self,
        file_path: str,
        property_id: str,
        language: str = "en"
    ) -> ParsedDocument:
        """
        Parse a local file.
        
        Args:
            file_path: Path to the file
            property_id: Property identifier
            language: Document language
            
        Returns:
            ParsedDocument with extracted content
        """
        filename = os.path.basename(file_path)
        
        with open(file_path, "rb") as f:
            content = f.read()
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            return await self._parse_pdf(content, "", filename, property_id, language)
        elif ext in [".html", ".htm"]:
            encoding = self.detect_encoding(content)
            text = content.decode(encoding, errors="replace")
            return self._parse_html(text, "", filename, property_id, language)
        elif ext in [".md", ".markdown"]:
            encoding = self.detect_encoding(content)
            text = content.decode(encoding, errors="replace")
            return self._parse_markdown(text, "", filename, property_id, language)
        elif ext in [".doc", ".docx"]:
            return await self._parse_docx(content, "", filename, property_id, language)
        elif ext == ".txt":
            encoding = self.detect_encoding(content)
            text = content.decode(encoding, errors="replace")
            return self._parse_plain_text(text, "", filename, property_id, language)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    async def _parse_pdf(
        self,
        content: bytes,
        url: str,
        filename: str,
        property_id: str,
        language: str
    ) -> ParsedDocument:
        """Parse PDF content."""
        try:
            from pypdf import PdfReader
            import io
            
            reader = PdfReader(io.BytesIO(content))
            
            full_text = ""
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n\n"
            
            # Extract sections from PDF text
            sections = self._extract_sections_from_text(full_text)
            doc_type = self._classify_document(full_text)
            
            title = filename or url.split("/")[-1] or "Untitled Document"
            if reader.metadata and reader.metadata.title:
                title = reader.metadata.title
            
            return ParsedDocument(
                title=title,
                sections=sections,
                doc_type=doc_type,
                language=language,
                source_url=url,
                source_filename=filename,
                raw_text=full_text,
                metadata={"pages": len(reader.pages)}
            )
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    async def _parse_docx(
        self,
        content: bytes,
        url: str,
        filename: str,
        property_id: str,
        language: str
    ) -> ParsedDocument:
        """Parse Word document."""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(content))
            
            sections = []
            current_section = None
            current_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if this is a heading
                if para.style.name.startswith("Heading"):
                    # Save previous section
                    if current_section is not None:
                        sections.append(ParsedSection(
                            title=current_section,
                            content="\n".join(current_content),
                            level=1,
                            is_critical=self._is_critical_section(current_section)
                        ))
                    current_section = text
                    current_content = []
                else:
                    current_content.append(text)
            
            # Don't forget last section
            if current_section is not None:
                sections.append(ParsedSection(
                    title=current_section,
                    content="\n".join(current_content),
                    level=1,
                    is_critical=self._is_critical_section(current_section)
                ))
            
            # If no sections found, create one from all text
            if not sections:
                full_text = "\n".join([p.text for p in doc.paragraphs])
                sections = self._extract_sections_from_text(full_text)
            
            full_text = "\n".join([p.text for p in doc.paragraphs])
            doc_type = self._classify_document(full_text)
            
            title = filename or url.split("/")[-1] or "Untitled Document"
            if doc.core_properties.title:
                title = doc.core_properties.title
            
            return ParsedDocument(
                title=title,
                sections=sections,
                doc_type=doc_type,
                language=language,
                source_url=url,
                source_filename=filename,
                raw_text=full_text,
                metadata={}
            )
        except Exception as e:
            raise ValueError(f"Failed to parse Word document: {str(e)}")
    
    def _parse_html(
        self,
        html: str,
        url: str,
        filename: str,
        property_id: str,
        language: str
    ) -> ParsedDocument:
        """Parse HTML content."""
        try:
            from bs4 import BeautifulSoup
            from markdownify import markdownify
            
            soup = BeautifulSoup(html, "html.parser")
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extract title
            title = filename or "Untitled"
            if soup.title:
                title = soup.title.string or title
            
            # Find main content
            main = soup.find("main") or soup.find("article") or soup.body or soup
            
            # Extract sections from headings
            sections = []
            headings = main.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
            
            for i, heading in enumerate(headings):
                heading_text = heading.get_text(strip=True)
                level = int(heading.name[1])
                
                # Get content until next heading
                content_parts = []
                for sibling in heading.find_next_siblings():
                    if sibling.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                        break
                    text = sibling.get_text(strip=True)
                    if text:
                        content_parts.append(text)
                
                if content_parts:
                    sections.append(ParsedSection(
                        title=heading_text,
                        content="\n".join(content_parts),
                        level=level,
                        is_critical=self._is_critical_section(heading_text)
                    ))
            
            # Get full text
            full_text = main.get_text(separator="\n", strip=True)
            
            # If no sections, extract from text
            if not sections:
                sections = self._extract_sections_from_text(full_text)
            
            doc_type = self._classify_document(full_text)
            
            return ParsedDocument(
                title=title,
                sections=sections,
                doc_type=doc_type,
                language=language,
                source_url=url,
                source_filename=filename,
                raw_text=full_text,
                metadata={}
            )
        except Exception as e:
            raise ValueError(f"Failed to parse HTML: {str(e)}")
    
    def _parse_markdown(
        self,
        text: str,
        url: str,
        filename: str,
        property_id: str,
        language: str
    ) -> ParsedDocument:
        """Parse Markdown content."""
        sections = []
        
        # Split by headings
        heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        # Find all headings
        matches = list(heading_pattern.finditer(text))
        
        for i, match in enumerate(matches):
            level = len(match.group(1))
            heading_text = match.group(2).strip()
            
            # Get content between this heading and next
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            
            # Remove markdown formatting from content
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)  # Bold
            content = re.sub(r'\*([^*]+)\*', r'\1', content)      # Italic
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)  # Links
            
            if content:
                sections.append(ParsedSection(
                    title=heading_text,
                    content=content,
                    level=level,
                    is_critical=self._is_critical_section(heading_text)
                ))
        
        if not sections:
            sections = self._extract_sections_from_text(text)
        
        doc_type = self._classify_document(text)
        title = filename or "Untitled"
        
        # Try to get title from first h1
        h1_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
        if h1_match:
            title = h1_match.group(1).strip()
        
        return ParsedDocument(
            title=title,
            sections=sections,
            doc_type=doc_type,
            language=language,
            source_url=url,
            source_filename=filename,
            raw_text=text,
            metadata={}
        )
    
    def _parse_plain_text(
        self,
        text: str,
        url: str,
        filename: str,
        property_id: str,
        language: str
    ) -> ParsedDocument:
        """Parse plain text content."""
        sections = self._extract_sections_from_text(text)
        doc_type = self._classify_document(text)
        title = filename or url.split("/")[-1] or "Untitled Document"
        
        return ParsedDocument(
            title=title,
            sections=sections,
            doc_type=doc_type,
            language=language,
            source_url=url,
            source_filename=filename,
            raw_text=text,
            metadata={}
        )
    
    def _extract_sections_from_text(self, text: str) -> List[ParsedSection]:
        """Extract sections from unstructured text using heuristics."""
        sections = []
        
        # Try to find section-like patterns (all caps lines, numbered sections, etc.)
        lines = text.split("\n")
        current_section = "General"
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect section headers (all caps, ends with colon, numbered)
            is_header = (
                (line.isupper() and len(line) > 3 and len(line) < 100) or
                (line.endswith(":") and len(line) < 80) or
                (re.match(r'^\d+[\.\)]\s+[A-Z]', line)) or
                (re.match(r'^[A-Z][A-Za-z\s]+:$', line))
            )
            
            if is_header and current_content:
                sections.append(ParsedSection(
                    title=current_section,
                    content="\n".join(current_content),
                    level=1,
                    is_critical=self._is_critical_section(current_section)
                ))
                current_section = line.rstrip(":")
                current_content = []
            else:
                current_content.append(line)
        
        # Add last section
        if current_content:
            sections.append(ParsedSection(
                title=current_section,
                content="\n".join(current_content),
                level=1,
                is_critical=self._is_critical_section(current_section)
            ))
        
        # If only one section with "General" title, break into paragraphs
        if len(sections) == 1 and sections[0].title == "General":
            paragraphs = text.split("\n\n")
            sections = []
            for i, para in enumerate(paragraphs):
                para = para.strip()
                if para and len(para) > 50:
                    first_line = para.split("\n")[0][:50]
                    sections.append(ParsedSection(
                        title=first_line + "..." if len(para.split("\n")[0]) > 50 else para.split("\n")[0],
                        content=para,
                        level=1,
                        is_critical=self._is_critical_section(para)
                    ))
        
        return sections if sections else [ParsedSection(
            title="Content",
            content=text,
            level=1,
            is_critical=False
        )]
    
    def _is_critical_section(self, text: str) -> bool:
        """Check if section contains critical information."""
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in self.CRITICAL_KEYWORDS)
    
    def _classify_document(self, text: str) -> DocType:
        """Classify document type based on content."""
        text_lower = text.lower()
        
        scores = {doc_type: 0 for doc_type in DocType}
        
        for doc_type, keywords in self.DOC_TYPE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    scores[doc_type] += 1
        
        # Get highest scoring type
        best_type = max(scores, key=scores.get)
        
        if scores[best_type] > 0:
            return best_type
        return DocType.UNKNOWN


# Singleton instance
_instance: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get or create the document parser singleton."""
    global _instance
    if _instance is None:
        _instance = DocumentParser()
    return _instance

